from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
styles = doc.styles

def set_style(style_name, font_name, font_size, bold=False, color=None):
    style = styles[style_name]
    style.font.name        = font_name
    style.font.size        = Pt(font_size)
    style.font.bold        = bold
    if color:
        style.font.color.rgb = RGBColor(*color)

set_style('Normal',   'Calibri', 11)
set_style('Heading 1','Calibri', 16, bold=True,  color=(0x1F, 0x45, 0x88))
set_style('Heading 2','Calibri', 13, bold=True,  color=(0x2E, 0x74, 0xB5))
set_style('Heading 3','Calibri', 11, bold=True,  color=(0x4A, 0x4A, 0x4A))

# ── Helper functions ──────────────────────────────────────────────────────────
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False, size=11, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.add_run(text)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue background
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2E74B5')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'),  'clear')
        tcPr.append(shd)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            row.cells[ci].text = cell_text
            row.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                tc   = row.cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'D6E4F7')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:val'),  'clear')
                tcPr.append(shd)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()   # spacing after table
    return table

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Light grey background via direct XML
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F2F2F2')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'),  'clear')
    pPr.append(shd)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("FAKE NEWS DETECTION SYSTEM")
r.font.size  = Pt(24)
r.font.bold  = True
r.font.color.rgb = RGBColor(0x1F, 0x45, 0x88)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Project Report — IBM Project")
r2.font.size   = Pt(14)
r2.font.italic = True
r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run("February 2026")
r3.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Table of Contents", 1)
toc_items = [
    ("1.", "Problem Statement"),
    ("2.", "Scope"),
    ("3.", "Goals / Objectives"),
    ("4.", "Background"),
    ("5.", "Feasibility Study"),
    ("6.", "Recommended Approach"),
    ("7.", "Recommended Tools & Technologies"),
    ("8.", "Doubts / Query"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(f"  {num}  {title}")
    r.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. Problem Statement", 1)
add_para(
    "The rapid spread of misinformation and fake news across digital platforms poses a significant "
    "threat to public knowledge, democratic processes, and social stability. Existing manual "
    "fact-checking methods are slow, resource-intensive, and unable to scale with the volume of "
    "online content published daily."
)
add_para(
    "At the same time, basic machine learning classifiers that rely on single signals (keyword "
    "matching or one model output) are easily fooled and suffer from keyword bias — for example, "
    "flagging legitimate journalism that simply reports about fake news."
)
add_para(
    "There is a pressing need for an accurate, explainable, and scalable automated system capable "
    "of classifying news articles as real or fake in real time — using a combination of deep "
    "learning and live external verification sources."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. SCOPE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. Scope", 1)
add_para("This project covers:", bold=True)
scope_items = [
    "AI Model: A Bidirectional LSTM deep learning model trained on 44,000+ labeled news articles (True.csv + Fake.csv)",
    "Hybrid Pipeline: A 4-layer scoring system combining the model's output with external APIs and heuristics",
    "Web Application: A Flask-based interface for users to submit text and receive verdicts with explanations",
    "User Management: Signup/login via MongoDB Atlas-backed accounts with magic-link email authentication",
    "External APIs: Google Fact Check Tools and NewsAPI for real-world corroboration of claims",
    "Explainability: Natural-language summaries telling the user why a result was given",
    "Resilience: Graceful error handling so the app never crashes when the database or APIs are unavailable",
]
for item in scope_items:
    add_bullet(item)

doc.add_paragraph()
add_para("Out of Scope:", bold=True)
out_of_scope = [
    "Video / audio fake news detection",
    "Multi-language support (currently English-only)",
    "Real-time social media scraping",
    "Per-user model fine-tuning",
]
for item in out_of_scope:
    add_bullet(item)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. GOALS / OBJECTIVES
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. Goals / Objectives", 1)
add_table(
    headers=["#", "Objective", "How Achieved"],
    rows=[
        ("1", "High classification accuracy",       "Bidirectional LSTM trained on 44K+ articles, 80/20 train-test split"),
        ("2", "Overcome keyword bias",               'Replace word "fake" → "fabricated" before tokenization'),
        ("3", "Corroborate with real-world data",   "Google Fact Check API and NewsAPI integrated as Layers 2 & 3"),
        ("4", "Explain every verdict",              "Human-readable summaries generated per detection method"),
        ("5", "Secure authentication",              "Magic-link tokens via itsdangerous, stored sessions"),
        ("6", "Resilient application",              "DB and API failures caught gracefully; app never shows crash page"),
        ("7", "Cloud-ready deployment",             "MongoDB Atlas for user data; model runs locally on Flask server"),
    ],
    col_widths=[1.2, 5.5, 8.3],
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. BACKGROUND
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4. Background", 1)

add_heading("4.1  The Problem Context", 2)
add_para(
    "Fake news causes measurable societal harm — from vaccine hesitancy driven by debunked claims "
    "to election interference. Manual fact-checking cannot keep pace with the volume of content "
    "published online daily. Automated AI-driven solutions are therefore essential."
)

add_heading("4.2  Why Deep Learning?", 2)
add_para(
    "Traditional bag-of-words or TF-IDF classifiers are brittle — they flag texts containing words "
    "like 'false' or 'fake' even when those words appear in legitimate investigative journalism. "
    "LSTM (Long Short-Term Memory) networks address this by understanding sequential context: "
    "the meaning of a sentence is built word-by-word, preserving order and relationships between words."
)

add_heading("4.3  Model Architecture", 2)
add_para("The project uses a Bidirectional LSTM — an upgraded LSTM that reads text both forward and backward, capturing richer context from both directions.")
add_code_block(
    "Input Text (max 300 tokens)\n"
    "        ↓\n"
    "Embedding Layer  (Vocab: 15,000 words  →  128-dimensional vectors)\n"
    "        ↓\n"
    "Bidirectional LSTM  (64 units,  return_sequences=True)\n"
    "        ↓\n"
    "Dropout (0.3)   ← prevents overfitting\n"
    "        ↓\n"
    "Bidirectional LSTM  (32 units)\n"
    "        ↓\n"
    "Dropout (0.3)\n"
    "        ↓\n"
    "Dense Layer  (64 neurons,  ReLU activation)\n"
    "        ↓\n"
    "Dropout (0.5)\n"
    "        ↓\n"
    "Output — Dense (1 neuron, Sigmoid)  →  probability 0.0 to 1.0\n"
    "         (> 0.55 = REAL,  < 0.45 = FAKE,  middle = UNCERTAIN)"
)

add_heading("4.4  Training Details", 2)
add_table(
    headers=["Parameter", "Value"],
    rows=[
        ("Dataset",                  "True.csv + Fake.csv  (~44,000+ articles total)"),
        ("Label encoding",           "Real = 1,  Fake = 0"),
        ("Vocabulary size",          "15,000 most frequent words"),
        ("Max sequence length",      "300 tokens (post-padded)"),
        ("Embedding dimensions",     "128"),
        ("Optimizer",                "Adam"),
        ("Loss function",            "Binary Cross-Entropy"),
        ("Train / Test split",       "80% / 20%  (random_state = 42)"),
        ("Batch size",               "64"),
        ("Max epochs",               "20  (with Early Stopping, patience = 2)"),
        ("Regularization",           "Dropout: 0.3 (LSTM layers),  0.5 (Dense layer)"),
        ("Model file size",          "~24 MB  (fake_news_model.h5)"),
        ("Tokenizer file size",      "~10 MB  (tokenizer.pkl)"),
    ],
    col_widths=[6, 9],
)

add_heading("4.5  Bias Mitigation", 2)
add_para(
    "A key innovation in this project: before tokenization, the word \"fake\" is replaced with "
    "\"fabricated\". This prevents the model from triggering a fake verdict simply because a "
    "legitimate article reports about fake news — a subtle but critical fix for reducing false positives."
)

add_heading("4.6  Training Outputs", 2)
for item in [
    "accuracy_graph.png — Training vs. validation accuracy per epoch",
    "loss_graph.png — Training vs. validation loss per epoch",
    "confusion_matrix.png — True/False positive/negative breakdown on the test set",
]:
    add_bullet(item)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. FEASIBILITY STUDY
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5. Feasibility Study", 1)
add_table(
    headers=["Dimension", "Assessment", "Notes"],
    rows=[
        ("Technical",        "✅ Feasible",   "TensorFlow, Flask, PyMongo ecosystem is mature and well-documented"),
        ("Data",             "✅ Feasible",   "44K+ labeled articles (Kaggle True/Fake CSVs) are publicly available"),
        ("Model Performance","✅ Feasible",   "BiLSTM on this dataset typically achieves 98–99% accuracy; confirmed by training graphs"),
        ("Operational",      "✅ Feasible",   "Trained model loads in <5 seconds; inference is near-instant per request"),
        ("Economic",         "✅ Low Cost",   "Free tiers: MongoDB Atlas, Google Fact Check API, NewsAPI cover development"),
        ("Time",             "✅ Feasible",   "Model training ~15–30 min on CPU; one-time cost, re-trainable manually"),
        ("Legal",            "✅ Compliant",  "No sensitive PII stored beyond email/name; passwords marked for hashing"),
        ("Risk",             "⚠️ Moderate",  "API rate limits (NewsAPI: 100 req/day on free tier); port 27017 blocked on some networks"),
    ],
    col_widths=[3.5, 3, 8.5],
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. RECOMMENDED APPROACH
# ══════════════════════════════════════════════════════════════════════════════
add_heading("6. Recommended Approach", 1)

add_heading("6.1  Model Pipeline (Core Engine)", 2)
add_code_block(
    "Raw User Text\n"
    "     │\n"
    "     ▼\n"
    " Text Cleaning\n"
    "  - Lowercase\n"
    '  - Replace "fake" → "fabricated"   ← bias mitigation\n'
    "  - Remove URLs (http / www)\n"
    "  - Remove non-alphabetic characters\n"
    "  - Remove English stopwords (NLTK)\n"
    "     │\n"
    "     ▼\n"
    " Keras Tokenizer  (vocab = 15,000)  →  integer sequences\n"
    "     │\n"
    "     ▼\n"
    " pad_sequences  →  fixed length 300\n"
    "     │\n"
    "     ▼\n"
    " Bidirectional LSTM Model  (fake_news_model.h5)\n"
    "     │\n"
    "     ▼\n"
    " raw_prob  (0.0 → 1.0,  sigmoid output)\n"
    "     │\n"
    "     ↓  feeds into the 4-Layer Hybrid Scoring System  ↓"
)

add_heading("6.2  4-Layer Hybrid Scoring System", 2)
add_para(
    "The raw model probability alone is not used as the final verdict. "
    "It is one input into a multi-layer hybrid system:"
)
add_table(
    headers=["Layer", "Signal", "Weight / Action"],
    rows=[
        ("Layer 1", "Known misinformation regex patterns\n(5G-COVID, vaccine-autism, flat earth, chemtrails, etc.)", "Hard override → FAKE"),
        ("Layer 2", "Google Fact Check API verdict\n(fact-checked FALSE / TRUE by Snopes, Reuters, AltNews, etc.)", "Strong override"),
        ("Layer 3", "NewsAPI — story found in trusted outlets?\n(BBC, NDTV, Reuters, AP, The Hindu, etc.)", "Adjusts hybrid score"),
        ("Layer 4", "Sensationalism heuristics + URL domain trust\n+ raw LSTM probability", "Fallback tiebreaker"),
    ],
    col_widths=[2.5, 7.5, 5],
)

add_para("Final thresholds:", bold=True)
for item in [
    "hybrid_prob > 0.55  →  REAL   (confidence = hybrid_prob × 100%)",
    "hybrid_prob < 0.45  →  FAKE   (confidence = (1 − hybrid_prob) × 100%)",
    "0.45 ≤ hybrid_prob ≤ 0.55  →  UNCERTAIN",
]:
    add_bullet(item)

add_heading("6.3  Additional Signals", 2)
for item in [
    "Sentiment (VADER): Positive / Negative / Neutral tone classification",
    "Subject classification: Politics, Health, Tech, Business, World News, General",
    "Source detection: URL count, citation count, 'according to' references",
]:
    add_bullet(item)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 7. RECOMMENDED TOOLS & TECHNOLOGIES
# ══════════════════════════════════════════════════════════════════════════════
add_heading("7. Recommended Tools & Technologies", 1)
add_table(
    headers=["Category", "Tool / Technology", "Specific Role in Project"],
    rows=[
        ("Deep Learning",        "TensorFlow 2.x + Keras",        "Build, train, save & load the BiLSTM model (.h5)"),
        ("NLP Preprocessing",   "NLTK",                          "Stopword removal (English), VADER sentiment scoring"),
        ("Tokenization",         "Keras Tokenizer",               "Maps words to integers, vocabulary size 15,000"),
        ("Sequence Handling",    "Keras pad_sequences",           "Pads/truncates input to max length 300"),
        ("Model Persistence",   "pickle (tokenizer), .h5 (model)","Save & reload trained model without retraining"),
        ("Backend Framework",    "Flask (Python)",                "Web server, REST API endpoints, session management"),
        ("Database",             "MongoDB Atlas (PyMongo)",       "User accounts: name, email, password, city"),
        ("Authentication",       "itsdangerous",                  "Time-limited magic-link tokens"),
        ("Email",                "smtplib + Gmail SMTP",          "Sends magic login links to users"),
        ("Fact-Checking API",    "Google Fact Check Tools",       "Layer 2 verdict corroboration from credible fact-checkers"),
        ("News Verification",   "NewsAPI",                        "Layer 3: checks if story appears in trusted outlets"),
        ("TLS Security",        "certifi, Python ssl",            "Secure MongoDB Atlas connections"),
        ("Data Processing",     "Pandas, NumPy",                 "Dataset loading and preprocessing during training"),
        ("Visualization",       "Matplotlib, Seaborn",           "Training accuracy/loss graphs, confusion matrix"),
        ("Model Evaluation",    "scikit-learn",                  "classification_report, accuracy_score, confusion_matrix"),
        ("Frontend",             "HTML5, CSS3, Jinja2",           "UI templates rendered by Flask"),
        ("Environment Config",  "python-dotenv",                 "Manages .env secrets: API keys, DB URI, SMTP creds"),
    ],
    col_widths=[4, 4.5, 6.5],
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. DOUBTS / QUERY
# ══════════════════════════════════════════════════════════════════════════════
add_heading("8. Doubts / Query", 1)
queries = [
    ("Model accuracy on Indian news",
     "The training data (True.csv, Fake.csv) is primarily sourced from US/Western news. "
     "How well does the model generalize to Indian news narratives (e.g., NDTV, The Hindu articles)?"),
    ("Re-training strategy",
     "The model is static after training. Should a retraining pipeline be scheduled (e.g., every 3 months) "
     "as new misinformation patterns emerge?"),
    ("UNCERTAIN zone handling",
     "When hybrid_prob falls between 0.45–0.55, the system returns 'Uncertain'. Should users be shown "
     "additional guidance or redirected to a manual fact-check link in this case?"),
    ("Raw model vs. hybrid fallback",
     "If Layer 1–3 APIs are unavailable (rate-limited or network blocked), the system falls back to "
     "raw LSTM + sensationalism heuristics. Is this fallback accuracy acceptable for production?"),
    ("Password security",
     "Passwords are currently stored in plain text in MongoDB. Should bcrypt hashing be added "
     "before this project is submitted or deployed publicly?"),
    ("Model file versioning",
     "The fake_news_model.h5 is ~24 MB and the tokenizer.pkl is ~10 MB. Should these be stored "
     "in cloud storage (IBM COS, AWS S3) instead of committed to the repository?"),
    ("Sequence length limitation",
     "Long articles are truncated to 300 tokens. Does this cause loss of important context "
     "for lengthy investigative pieces?"),
]
for i, (title, body) in enumerate(queries, start=1):
    p = doc.add_paragraph()
    run = p.add_run(f"Q{i}.  {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x45, 0x88)
    add_para(body, indent=True)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = r"c:\Users\Trishla\Desktop\IBM Project\Fake News Detection\FakeNewsDetection_ProjectReport.docx"
doc.save(output_path)
print(f"SUCCESS: Report saved to {output_path}")
